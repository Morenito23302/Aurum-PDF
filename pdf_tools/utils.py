import os
import json as _json
import fitz  # PyMuPDF
import tempfile
import zipfile
import pdfplumber
import pandas as pd
from io import BytesIO
from pdf2image import convert_from_path
import pytesseract
from docx import Document
import subprocess
from PIL import Image

def merge_pdfs_util(file_paths, output_path):
    # PyMuPDF para unir
    result_pdf = fitz.open()
    for fp in file_paths:
        with fitz.open(fp) as doc:
            result_pdf.insert_pdf(doc)
    result_pdf.save(output_path)
    result_pdf.close()

def is_pdf_scanned(pdf_path):
    """
    Detecta si un PDF es escaneado (sin texto digital) o digital.
    Devuelve True si es escaneado, False si tiene texto.
    """
    try:
        doc = fitz.open(pdf_path)
        for page in doc:
            if page.get_text("text").strip():
                doc.close()
                return False
        doc.close()
        return True
    except Exception:
        return True

def ocr_pdf_to_word(pdf_path, output_path):
    """
    Convierte un PDF escaneado a Word usando OCR página por página de forma eficiente.
    Optimizado para servidores con poca RAM (como Render Free).
    """
    from docx.shared import Inches
    from pdf2image import convert_from_path, get_page_count
    
    doc_docx = Document()
    
    # Configurar márgenes
    sections = doc_docx.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    try:
        # Verificar que tesseract esté instalado
        import shutil
        if not shutil.which("tesseract"):
            raise Exception("El motor de OCR (Tesseract) no está instalado en el servidor.")

        page_count = get_page_count(pdf_path)
        print(f"Iniciando OCR de {page_count} páginas...")

        for i in range(1, page_count + 1):
            # Convertir una sola página a la vez para ahorrar RAM
            # Bajamos a 200 DPI (equilibrio entre velocidad/memoria y precisión)
            pages = convert_from_path(pdf_path, dpi=200, first_page=i, last_page=i)
            if not pages:
                continue
            
            image = pages[0]
            
            # 1. Ejecutar OCR
            text = pytesseract.image_to_string(image, lang='spa+eng')
            
            # 2. Añadir texto
            if text.strip():
                doc_docx.add_paragraph(text)
            else:
                # Si no hay texto, poner un aviso o dejar espacio
                doc_docx.add_paragraph(f"[Página {i} sin texto detectable]")
            
            # 3. Salto de página
            if i < page_count:
                doc_docx.add_page_break()
            
            # Liberar memoria de la imagen explícitamente
            del image
            del pages
                
        doc_docx.save(output_path)
    except Exception as e:
        print(f"Error en ocr_pdf_to_word: {str(e)}")
        raise Exception(f"Fallo en el motor de OCR: {str(e)}")

def convert_to_word_util(pdf_path, output_path, mode='auto'):
    """
    Convierte un PDF a DOCX con una cadena de fallbacks ultra-resistente.
    1. OCR (si es necesario o solicitado)
    2. pdf2docx (Mejor layout)
    3. LibreOffice (Máxima compatibilidad)
    """
    import shutil
    from pdf2docx import Converter

    # 1. ¿Usamos OCR?
    use_ocr = (mode == 'ocr') or (mode == 'auto' and is_pdf_scanned(pdf_path))

    if use_ocr:
        print(f"--- Iniciando modo OCR para: {pdf_path} ---")
        try:
            return ocr_pdf_to_word(pdf_path, output_path)
        except Exception as e:
            print(f"OCR falló: {e}. Intentando métodos digitales por si acaso...")

    # 2. Intentar pdf2docx (Estrategia principal para digital)
    print(f"--- Intentando pdf2docx para: {pdf_path} ---")
    try:
        cv = Converter(pdf_path)
        # kwargs cpu_count=1 para evitar usar demasiada RAM creando procesos
        cv.convert(output_path, start=0, end=None, cpu_count=1)
        cv.close()
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            return
    except Exception as e:
        print(f"pdf2docx falló críticamente: {e}")
        # Limpiar si dejó archivo corrupto
        if os.path.exists(output_path):
            os.remove(output_path)

    # 3. Intentar LibreOffice (Último recurso, muy estable)
    print(f"--- Intentando LibreOffice para: {pdf_path} ---")
    soffice_path = shutil.which("libreoffice") or shutil.which("soffice")
    if soffice_path:
        try:
            output_dir = os.path.dirname(output_path)
            # Usar un HOME temporal para LibreOffice
            env = os.environ.copy()
            env['HOME'] = '/tmp'
            
            subprocess.run(
                [soffice_path, "--headless", "--infilter=writer_pdf_import",
                 "--convert-to", "docx:MS Word 2007 XML", pdf_path, "--outdir", output_dir],
                check=True, timeout=180, env=env
            )
            base_name = os.path.splitext(os.path.basename(pdf_path))[0]
            generated = os.path.join(output_dir, base_name + ".docx")
            if os.path.exists(generated):
                if generated != output_path:
                    os.replace(generated, output_path)
                return
        except Exception as e:
            print(f"LibreOffice falló críticamente: {e}")

    # 4. Si todo falla, intentar OCR si no se intentó antes
    if not use_ocr:
        print("Todo lo digital falló. Intentando OCR como último recurso...")
        return ocr_pdf_to_word(pdf_path, output_path)

    raise Exception("Lo sentimos, no pudimos procesar este PDF. El archivo podría estar protegido o ser incompatible con los conversores actuales.")


def convert_to_excel_util(pdf_path, output_path):
    # Usar pdfplumber y pandas para volcar texto o tablas a excel
    with pdfplumber.open(pdf_path) as pdf:
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            data_found = False
            for i, page in enumerate(pdf.pages):
                tables = page.extract_tables()
                if tables:
                    for j, table in enumerate(tables):
                        if not table:
                            continue
                        data_found = True
                        df = pd.DataFrame(table[1:], columns=table[0])
                        # Limpiar encabezados si son nulos
                        if len(df.columns) > 0:
                            df.columns = [str(col) if col else f"Col_{c_idx}" for c_idx, col in enumerate(df.columns)]
                        sheet_name = f"Pag{i+1}_Tab{j+1}"[:31]
                        df.to_excel(writer, sheet_name=sheet_name, index=False)
                else:
                    # Extraer texto si no hay tablas
                    text = page.extract_text()
                    if text and text.strip():
                        data_found = True
                        # Tratar de separar en columnas por espacios múltiples o simplemente dejar líneas
                        lines = [line.split() for line in text.split('\n') if line.strip()]
                        if lines:
                            df = pd.DataFrame(lines)
                            sheet_name = f"Pag{i+1}_Texto"[:31]
                            df.to_excel(writer, sheet_name=sheet_name, index=False, header=False)
                
                # ¡MUY IMPORTANTE PARA EVITAR OOM ERROR (502 BAD GATEWAY) EN SERVIDORES GRATIS!
                page.flush_cache()
            
            if not data_found:
                # Escribir hoja vacía si no hay nada
                pd.DataFrame([["No se encontraron datos tabulares o texto"]]).to_excel(writer, index=False, header=False)

def extract_images_util(pdf_path, zip_output_path):
    pdf_document = fitz.open(pdf_path)
    with zipfile.ZipFile(zip_output_path, 'w') as zipf:
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = pdf_document.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                image_filename = f"image_page{page_num+1}_{img_index+1}.{image_ext}"
                zipf.writestr(image_filename, image_bytes)
    pdf_document.close()

def convert_to_pdf_util(input_paths, output_path):
    """
    Convierte uno o varios archivos (imágenes o documentos) a un único PDF.
    """
    temp_pdfs = []
    temp_dir = tempfile.mkdtemp()
    
    try:
        for idx, input_path in enumerate(input_paths):
            ext = os.path.splitext(input_path)[1].lower()
            temp_out = os.path.join(temp_dir, f"part_{idx}.pdf")
            
            if ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
                image = Image.open(input_path)
                if image.mode != 'RGB':
                    image = image.convert('RGB')
                image.save(temp_out, "PDF", resolution=100.0)
                temp_pdfs.append(temp_out)
            elif ext in ['.docx', '.doc', '.xlsx', '.xls', '.ppt', '.pptx']:
                try:
                    subprocess.run([
                        'libreoffice', '--headless', '--convert-to', 'pdf',
                        input_path, '--outdir', temp_dir
                    ], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                    
                    base_name = os.path.splitext(os.path.basename(input_path))[0]
                    generated_pdf = os.path.join(temp_dir, base_name + '.pdf')
                    
                    if os.path.exists(generated_pdf):
                        os.rename(generated_pdf, temp_out)
                        temp_pdfs.append(temp_out)
                except Exception as e:
                    print(f"Error convirtiendo {input_path} con LibreOffice: {e}")
            elif ext == '.pdf':
                # Si ya es PDF, solo lo añadimos para unir
                temp_pdfs.append(input_path)

        if not temp_pdfs:
            raise Exception("No se pudo convertir ninguno de los archivos seleccionados.")

        # Unir todos los PDFs generados
        merge_pdfs_util(temp_pdfs, output_path)
        
    finally:
        # Limpieza (opcional aquí si tempfile maneja el ciclo, pero mejor ser explícito)
        pass


def extract_text_blocks_util(pdf_path):
    """Extrae todos los bloques de texto (spans) del PDF con posición, fuente y color."""
    doc = fitz.open(pdf_path)
    result = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        page_dict = page.get_text(
            "dict",
            flags=fitz.TEXT_PRESERVE_WHITESPACE | fitz.TEXT_PRESERVE_LIGATURES
        )
        page_data = {
            "page": page_num,
            "width": round(page.rect.width, 2),
            "height": round(page.rect.height, 2),
            "blocks": []
        }
        counter = 0
        for block in page_dict.get("blocks", []):
            if block.get("type") != 0:   # 0 = text, 1 = image
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "")
                    if not text.strip():
                        continue
                    color_int = span.get("color", 0)
                    r = (color_int >> 16) & 0xFF
                    g = (color_int >> 8) & 0xFF
                    b = color_int & 0xFF
                    bbox = span["bbox"]
                    origin = span.get("origin", (bbox[0], bbox[3]))
                    page_data["blocks"].append({
                        "id": f"p{page_num}_s{counter}",
                        "text": text,
                        "x0": round(bbox[0], 2),
                        "y0": round(bbox[1], 2),
                        "x1": round(bbox[2], 2),
                        "y1": round(bbox[3], 2),
                        "origin_y": round(origin[1], 2),  # baseline exacto
                        "size": round(span.get("size", 12), 2),
                        "font": span.get("font", "Helvetica"),
                        "color_hex": "#{:02x}{:02x}{:02x}".format(r, g, b),
                        "flags": span.get("flags", 0),
                    })
                    counter += 1
        result.append(page_data)
    doc.close()
    return result


def _map_font(font_name, flags=0):
    """Mapea nombre de fuente del PDF a una de las fuentes base-14 de PyMuPDF."""
    # Quitar prefijo de subset p.ej. "ABCDEF+ArialMT" → "ArialMT"
    name = font_name.split("+")[-1].lower()
    is_bold   = bool(flags & 16) or "bold" in name or "black" in name or "heavy" in name
    is_italic = bool(flags & 2)  or "italic" in name or "oblique" in name or "it" in name

    if any(x in name for x in ["helvetica", "arial", "calibri", "tahoma",
                                 "verdana", "gothic", "futura", "gill"]):
        if is_bold and is_italic: return "helv"   # no helv-bi in base14, use helv
        if is_bold:               return "hebo"
        if is_italic:             return "heoi"
        return "helv"
    elif any(x in name for x in ["times", "palatino", "garamond", "georgia",
                                  "roman", "minion", "caslon"]):
        if is_bold and is_italic: return "tibi"
        if is_bold:               return "tibo"
        if is_italic:             return "tiit"
        return "tiro"
    elif any(x in name for x in ["courier", "mono", "consolas", "menlo",
                                   "inconsolata", "lucidacon"]):
        if is_bold and is_italic: return "cobi"
        if is_bold:               return "cobo"
        if is_italic:             return "coit"
        return "cour"
    return "helv"   # fallback universal


def apply_text_edits_util(pdf_path, output_path, edits):
    """
    Aplica ediciones de texto al PDF:
    - Redacta (borra) el texto original con un rectángulo del color de fondo.
    - Re-inserta el nuevo texto en la misma posición.
    Imágenes, tablas, encabezados, logos y marcas de agua NO se tocan.

    edits: lista de dicts con claves:
        page (int), x0/y0/x1/y1 (float), font (str), flags (int),
        new_text (str), color_hex (str), size (float)
    """
    from collections import defaultdict
    doc = fitz.open(pdf_path)

    page_edits = defaultdict(list)
    for edit in edits:
        page_edits[int(edit["page"])].append(edit)

    for page_num, edit_list in page_edits.items():
        if page_num >= len(doc):
            continue
        page = doc[page_num]

        # --- Paso 1: marcar redacciones QUIRÚRGICAS (solo el área del texto) ---
        for edit in edit_list:
            # En lugar de usar todo el bbox (que suele tocar líneas), 
            # usamos un área más pequeña centrada en el baseline.
            base_y = edit.get("origin_y", edit["y1"])
            size = edit.get("size", 12)
            # Altura quirúrgica: de baseline - 80% de size hasta baseline + 20%
            rect = fitz.Rect(edit["x0"], base_y - (size * 0.8), edit["x1"], base_y + (size * 0.2))
            
            # Marcamos redacción con fondo blanco
            page.add_redact_annot(rect, fill=(1, 1, 1)) 
        
        # Aplicamos redacciones PRESERVANDO estrictamente gráficos (líneas de tablas)
        page.apply_redactions(
            images=fitz.PDF_REDACT_IMAGE_NONE,
            graphics=fitz.PDF_REDACT_LINE_ART_NONE
        )

        # --- Paso 2: re-insertar texto ---
        for edit in edit_list:
            new_text = edit.get("new_text", "").strip()
            if not new_text:
                continue   # bloque eliminado intencionalmente

            color_hex = edit.get("color_hex", "#000000").lstrip("#")
            try:
                r = int(color_hex[0:2], 16) / 255.0
                g = int(color_hex[2:4], 16) / 255.0
                b = int(color_hex[4:6], 16) / 255.0
            except Exception:
                r, g, b = 0.0, 0.0, 0.0

            fontname = _map_font(edit.get("font", "Helvetica"), edit.get("flags", 0))
            size     = float(edit.get("size", 12))

            # Usar origin_y (baseline exacto) en lugar de y1 para evitar corrimiento vertical
            baseline_y = edit.get("origin_y", edit["y1"])
            page.insert_text(
                fitz.Point(edit["x0"], baseline_y),
                new_text,
                fontsize=size,
                fontname=fontname,
                color=(r, g, b),
            )

    doc.save(output_path, garbage=4, deflate=True, clean=True)
    doc.close()


def protect_unlock_pdf_util(pdf_path, output_path, password="", mode="unlock"):
    """Elimina contraseñas y restricciones, o añade una contraseña al PDF usando PyMuPDF."""
    doc = fitz.open(pdf_path)
    if mode == "unlock":
        if doc.is_encrypted:
            if not doc.authenticate(password):
                doc.close()
                raise Exception("La contraseña es incorrecta o el documento está fuertemente cifrado y requiere una contraseña de apertura válida.")
        doc.save(output_path)
    else:  # mode == "lock"
        if not password:
            doc.close()
            raise Exception("Se requiere una contraseña para proteger el PDF.")
        if doc.is_encrypted:
            if not doc.authenticate(""):
                doc.authenticate(password)
        doc.save(output_path, encryption=fitz.PDF_ENCRYPT_AES_256, owner_pw=password, user_pw=password)
    doc.close()
